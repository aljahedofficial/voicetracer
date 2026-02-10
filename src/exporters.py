"""
VoiceTracer Export Module

Generates exportable reports in multiple formats (PDF, DOCX, XLSX, PPTX, CSV, JSON, PNG, ZIP).
"""

import json
import csv
import io
from pathlib import Path
from typing import Dict, List, Any, BinaryIO
from datetime import datetime
import pandas as pd
import plotly.io as pio
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib import colors
from metrics_spec import MetricType, normalize_metric
from visualizations import (
    RadarChartGenerator,
    BarChartGenerator,
    DeltaVisualization,
    BurstinessVisualization,
    IndividualMetricCharts,
)


class ExportMetadata:
    """Generate metadata for exports."""
    
    @staticmethod
    def create_metadata(
        analysis_result,
        doc_pair,
        original_metadata: Dict,
        edited_metadata: Dict,
        calibration: Dict = None,
    ) -> Dict[str, Any]:
        """Create metadata for export."""
        metadata = {
            'title': 'VoiceTracer Analysis Report',
            'created_at': datetime.now().isoformat(),
            'analysis_id': analysis_result.doc_pair_id,
            'original_text_stats': original_metadata,
            'edited_text_stats': edited_metadata,
            'metrics': {
                'original': {
                    'burstiness': analysis_result.original_metrics.burstiness,
                    'lexical_diversity': analysis_result.original_metrics.lexical_diversity,
                    'syntactic_complexity': analysis_result.original_metrics.syntactic_complexity,
                    'ai_ism_likelihood': analysis_result.original_metrics.ai_ism_likelihood,
                    'function_word_ratio': analysis_result.original_metrics.function_word_ratio,
                    'discourse_marker_density': analysis_result.original_metrics.discourse_marker_density,
                    'information_density': analysis_result.original_metrics.information_density,
                    'epistemic_hedging': analysis_result.original_metrics.epistemic_hedging,
                },
                'edited': {
                    'burstiness': analysis_result.edited_metrics.burstiness,
                    'lexical_diversity': analysis_result.edited_metrics.lexical_diversity,
                    'syntactic_complexity': analysis_result.edited_metrics.syntactic_complexity,
                    'ai_ism_likelihood': analysis_result.edited_metrics.ai_ism_likelihood,
                    'function_word_ratio': analysis_result.edited_metrics.function_word_ratio,
                    'discourse_marker_density': analysis_result.edited_metrics.discourse_marker_density,
                    'information_density': analysis_result.edited_metrics.information_density,
                    'epistemic_hedging': analysis_result.edited_metrics.epistemic_hedging,
                },
                'deltas': {
                    'burstiness_delta': analysis_result.metric_deltas.burstiness_delta,
                    'burstiness_pct_change': analysis_result.metric_deltas.burstiness_pct_change,
                    'lexical_diversity_delta': analysis_result.metric_deltas.lexical_diversity_delta,
                    'lexical_diversity_pct_change': analysis_result.metric_deltas.lexical_diversity_pct_change,
                    'syntactic_complexity_delta': analysis_result.metric_deltas.syntactic_complexity_delta,
                    'syntactic_complexity_pct_change': analysis_result.metric_deltas.syntactic_complexity_pct_change,
                    'ai_ism_delta': analysis_result.metric_deltas.ai_ism_delta,
                    'ai_ism_pct_change': analysis_result.metric_deltas.ai_ism_pct_change,
                    'function_word_ratio_delta': analysis_result.metric_deltas.function_word_ratio_delta,
                    'function_word_ratio_pct_change': analysis_result.metric_deltas.function_word_ratio_pct_change,
                    'discourse_marker_density_delta': analysis_result.metric_deltas.discourse_marker_density_delta,
                    'discourse_marker_density_pct_change': analysis_result.metric_deltas.discourse_marker_density_pct_change,
                    'information_density_delta': analysis_result.metric_deltas.information_density_delta,
                    'information_density_pct_change': analysis_result.metric_deltas.information_density_pct_change,
                    'epistemic_hedging_delta': analysis_result.metric_deltas.epistemic_hedging_delta,
                    'epistemic_hedging_pct_change': analysis_result.metric_deltas.epistemic_hedging_pct_change,
                },
            },
            'thesis_alignment': {
                'research_questions_addressed': [
                    'How does AI editing affect L2 learner writing characteristics?',
                    'Can we quantify stylistic differences between human and AI-edited text?',
                    'What linguistic markers indicate AI involvement?',
                    'How can L2 learners preserve voice while improving grammar?',
                ],
                'methodology': 'VoiceTracer v0.1.0',
                'source': 'https://github.com/VoiceTracer',
            },
        }

        if calibration:
            metadata['calibration'] = calibration

        return metadata


def _score_against_standards(value: float, human: float, ai: float) -> float:
    if human == ai:
        return 0.5
    score = (value - ai) / (human - ai)
    return max(0.0, min(score, 1.0))


def _position_against_standards(value: float, human: float, ai: float) -> float:
    if human == ai:
        return 0.5
    position = (value - human) / (ai - human)
    return max(0.0, min(position, 1.0))


def _verdict_from_shift(shift: float) -> str:
    if abs(shift) < 0.15:
        return "Preserved (minimal change)"
    if shift > 0.30:
        return "Compromised (significant homogenization)"
    if shift > 0.15:
        return "Moderate shift (partial homogenization)"
    return "Enhanced (toward human norm)"


def _build_metric_verdicts(analysis_result, calibration: Dict = None) -> Dict[str, str]:
    verdicts = {}
    calibration = calibration or {}
    adjusted = calibration.get("adjusted", calibration.get("default", {}))
    human = adjusted.get("human", {})
    ai = adjusted.get("ai", {})

    metric_keys = [
        "burstiness",
        "lexical_diversity",
        "syntactic_complexity",
        "ai_ism_likelihood",
        "function_word_ratio",
        "discourse_marker_density",
        "information_density",
        "epistemic_hedging",
    ]

    for key in metric_keys:
        orig_val = getattr(analysis_result.original_metrics, key)
        edit_val = getattr(analysis_result.edited_metrics, key)
        human_val = human.get(key, 0.0)
        ai_val = ai.get(key, 0.0)

        orig_pos = _position_against_standards(orig_val, human_val, ai_val)
        edit_pos = _position_against_standards(edit_val, human_val, ai_val)
        verdicts[key] = _verdict_from_shift(edit_pos - orig_pos)

    return verdicts


def _normalized_metric_dict(metric_scores) -> Dict[str, float]:
    return {
        "burstiness": normalize_metric(metric_scores.burstiness, MetricType.BURSTINESS),
        "lexical_diversity": normalize_metric(metric_scores.lexical_diversity, MetricType.LEXICAL_DIVERSITY),
        "syntactic_complexity": normalize_metric(metric_scores.syntactic_complexity, MetricType.SYNTACTIC_COMPLEXITY),
        "ai_ism_likelihood": normalize_metric(metric_scores.ai_ism_likelihood, MetricType.AI_ISM_LIKELIHOOD),
        "function_word_ratio": normalize_metric(metric_scores.function_word_ratio, MetricType.FUNCTION_WORD_RATIO),
        "discourse_marker_density": normalize_metric(metric_scores.discourse_marker_density, MetricType.DISCOURSE_MARKER_DENSITY),
        "information_density": normalize_metric(metric_scores.information_density, MetricType.INFORMATION_DENSITY),
        "epistemic_hedging": normalize_metric(metric_scores.epistemic_hedging, MetricType.EPISTEMIC_HEDGING),
    }


def _raw_metric_dict(metric_scores) -> Dict[str, float]:
    return {
        "burstiness_raw": metric_scores.burstiness,
        "lexical_diversity_raw": metric_scores.lexical_diversity,
        "syntactic_complexity_raw": metric_scores.syntactic_complexity,
        "ai_ism_likelihood_raw": metric_scores.ai_ism_likelihood,
        "function_word_ratio_raw": metric_scores.function_word_ratio,
        "discourse_marker_density_raw": metric_scores.discourse_marker_density,
        "information_density_raw": metric_scores.information_density,
        "epistemic_hedging_raw": metric_scores.epistemic_hedging,
    }


def _plotly_to_image(fig, width: int = 1200, height: int = 800, scale: int = 2) -> bytes:
    try:
        return pio.to_image(
            fig,
            format="png",
            width=width,
            height=height,
            scale=scale,
            engine="kaleido",
        )
    except Exception:
        return fig.to_image(format="png", width=width, height=height, scale=scale)


class CSVExporter:
    """Export data to CSV format."""
    
    @staticmethod
    def export(
        analysis_result,
        doc_pair,
        original_metadata: Dict,
        edited_metadata: Dict,
        include_original_text: bool = False,
        include_edited_text: bool = False,
        **options
    ) -> str:
        """
        Generate CSV export with analysis data.
        
        Returns:
            CSV string
        """
        output = io.StringIO()
        writer = csv.writer(output)

        # Header section
        writer.writerow(['VoiceTracer Analysis Report'])
        writer.writerow(['Generated:', datetime.now().isoformat()])
        writer.writerow([])

        # Text statistics
        writer.writerow(['Text Statistics'])
        writer.writerow(['Metric', 'Original', 'Edited', 'Delta', 'Absolute Shift (Δ)'])
        writer.writerow(['Word Count',
                        original_metadata.get('word_count', 0),
                        edited_metadata.get('word_count', 0),
                        edited_metadata.get('word_count', 0) - original_metadata.get('word_count', 0),
                        ''])
        writer.writerow(['Character Count',
                        original_metadata.get('char_count', 0),
                        edited_metadata.get('char_count', 0),
                        edited_metadata.get('char_count', 0) - original_metadata.get('char_count', 0),
                        ''])
        writer.writerow(['Sentence Count',
                        original_metadata.get('sentence_count', 0),
                        edited_metadata.get('sentence_count', 0),
                        edited_metadata.get('sentence_count', 0) - original_metadata.get('sentence_count', 0),
                        ''])
        writer.writerow([])

        # Metrics
        writer.writerow(['Metric Comparison'])
        writer.writerow(['Metric', 'Original', 'Edited', 'Delta', 'Absolute Shift (Δ)'])

        metrics_to_export = [
            ('Burstiness', 'burstiness'),
            ('Lexical Diversity', 'lexical_diversity'),
            ('Syntactic Complexity', 'syntactic_complexity'),
            ('AI-ism Likelihood', 'ai_ism_likelihood'),
            ('Function Word Ratio', 'function_word_ratio'),
            ('Discourse Marker Density', 'discourse_marker_density'),
            ('Information Density', 'information_density'),
            ('Epistemic Hedging', 'epistemic_hedging'),
        ]
        metrics_delta_map = {
            'burstiness': (analysis_result.metric_deltas.burstiness_delta,
                           analysis_result.metric_deltas.burstiness_pct_change),
            'lexical_diversity': (analysis_result.metric_deltas.lexical_diversity_delta,
                                  analysis_result.metric_deltas.lexical_diversity_pct_change),
            'syntactic_complexity': (analysis_result.metric_deltas.syntactic_complexity_delta,
                                     analysis_result.metric_deltas.syntactic_complexity_pct_change),
            'ai_ism_likelihood': (analysis_result.metric_deltas.ai_ism_delta,
                                  analysis_result.metric_deltas.ai_ism_pct_change),
            'function_word_ratio': (analysis_result.metric_deltas.function_word_ratio_delta,
                                    analysis_result.metric_deltas.function_word_ratio_pct_change),
            'discourse_marker_density': (analysis_result.metric_deltas.discourse_marker_density_delta,
                                         analysis_result.metric_deltas.discourse_marker_density_pct_change),
            'information_density': (analysis_result.metric_deltas.information_density_delta,
                                    analysis_result.metric_deltas.information_density_pct_change),
            'epistemic_hedging': (analysis_result.metric_deltas.epistemic_hedging_delta,
                                  analysis_result.metric_deltas.epistemic_hedging_pct_change),
        }

        for label, key in metrics_to_export:
            orig = getattr(analysis_result.original_metrics, key)
            edit = getattr(analysis_result.edited_metrics, key)
            delta, pct = metrics_delta_map.get(key, (edit - orig, 0))

            writer.writerow([label, round(orig, 3), round(edit, 3), round(delta, 3), f'{pct:+.1f}%'])

        writer.writerow([])

        # AI-isms detected
        if analysis_result.ai_isms:
            writer.writerow(['AI-isms Detected'])
            writer.writerow(['Phrase', 'Category', 'Context'])
            for ai_ism in analysis_result.ai_isms:
                writer.writerow([
                    ai_ism.get('phrase', ''),
                    ai_ism.get('category', ''),
                    ai_ism.get('context', '')[:100],  # Truncate context
                ])

        calibration = options.get('calibration')
        if calibration:
            writer.writerow([])
            writer.writerow(['Calibration Standards'])
            writer.writerow(['Mode', 'Metric', 'Human Standard', 'AI Standard'])
            for mode in ['default', 'adjusted']:
                for metric, human_value in calibration.get(mode, {}).get('human', {}).items():
                    ai_value = calibration.get(mode, {}).get('ai', {}).get(metric, '')
                    writer.writerow([mode, metric, human_value, ai_value])

            writer.writerow([])
            writer.writerow(['Calibration Impact (Score Δ)'])
            writer.writerow(['Metric', 'Original Δ', 'Edited Δ'])
            impact = calibration.get('impact', {})
            for metric, orig_delta in impact.get('original', {}).items():
                edit_delta = impact.get('edited', {}).get(metric, '')
                writer.writerow([metric, round(orig_delta, 3), round(edit_delta, 3)])

            notes = calibration.get('notes', {})
            if notes:
                writer.writerow([])
                writer.writerow(['Calibration Notes'])
                for key, text in notes.items():
                    writer.writerow([key, text])

        return output.getvalue()


class JSONExporter:
    """Export data to JSON format."""
    
    @staticmethod
    def export(
        analysis_result,
        doc_pair,
        original_metadata: Dict,
        edited_metadata: Dict,
        include_original_text: bool = False,
        include_edited_text: bool = False,
        **options
    ) -> str:
        """
        Generate JSON export with full analysis data.
        
        Returns:
            JSON string
        """
        export_data = {
            'metadata': ExportMetadata.create_metadata(
                analysis_result,
                doc_pair,
                original_metadata,
                edited_metadata,
                calibration=options.get('calibration'),
            ),
            'text_statistics': {
                'original': original_metadata,
                'edited': edited_metadata,
            },
            'metrics': {
                'original': {
                    'burstiness': analysis_result.original_metrics.burstiness,
                    'lexical_diversity': analysis_result.original_metrics.lexical_diversity,
                    'syntactic_complexity': analysis_result.original_metrics.syntactic_complexity,
                    'ai_ism_likelihood': analysis_result.original_metrics.ai_ism_likelihood,
                    'function_word_ratio': analysis_result.original_metrics.function_word_ratio,
                    'discourse_marker_density': analysis_result.original_metrics.discourse_marker_density,
                    'information_density': analysis_result.original_metrics.information_density,
                    'epistemic_hedging': analysis_result.original_metrics.epistemic_hedging,
                },
                'edited': {
                    'burstiness': analysis_result.edited_metrics.burstiness,
                    'lexical_diversity': analysis_result.edited_metrics.lexical_diversity,
                    'syntactic_complexity': analysis_result.edited_metrics.syntactic_complexity,
                    'ai_ism_likelihood': analysis_result.edited_metrics.ai_ism_likelihood,
                    'function_word_ratio': analysis_result.edited_metrics.function_word_ratio,
                    'discourse_marker_density': analysis_result.edited_metrics.discourse_marker_density,
                    'information_density': analysis_result.edited_metrics.information_density,
                    'epistemic_hedging': analysis_result.edited_metrics.epistemic_hedging,
                },
                'deltas': {
                    'burstiness': {
                        'delta': analysis_result.metric_deltas.burstiness_delta,
                        'pct_change': analysis_result.metric_deltas.burstiness_pct_change,
                    },
                    'lexical_diversity': {
                        'delta': analysis_result.metric_deltas.lexical_diversity_delta,
                        'pct_change': analysis_result.metric_deltas.lexical_diversity_pct_change,
                    },
                    'syntactic_complexity': {
                        'delta': analysis_result.metric_deltas.syntactic_complexity_delta,
                        'pct_change': analysis_result.metric_deltas.syntactic_complexity_pct_change,
                    },
                    'ai_ism_likelihood': {
                        'delta': analysis_result.metric_deltas.ai_ism_delta,
                        'pct_change': analysis_result.metric_deltas.ai_ism_pct_change,
                    },
                    'function_word_ratio': {
                        'delta': analysis_result.metric_deltas.function_word_ratio_delta,
                        'pct_change': analysis_result.metric_deltas.function_word_ratio_pct_change,
                    },
                    'discourse_marker_density': {
                        'delta': analysis_result.metric_deltas.discourse_marker_density_delta,
                        'pct_change': analysis_result.metric_deltas.discourse_marker_density_pct_change,
                    },
                    'information_density': {
                        'delta': analysis_result.metric_deltas.information_density_delta,
                        'pct_change': analysis_result.metric_deltas.information_density_pct_change,
                    },
                    'epistemic_hedging': {
                        'delta': analysis_result.metric_deltas.epistemic_hedging_delta,
                        'pct_change': analysis_result.metric_deltas.epistemic_hedging_pct_change,
                    },
                },
            },
            'ai_isms_detected': [
                {
                    'phrase': ai.get('phrase'),
                    'category': ai.get('category'),
                    'context': ai.get('context'),
                }
                for ai in analysis_result.ai_isms
            ],
        }
        
        if include_original_text:
            export_data['texts'] = {'original': doc_pair.original_text}
        if include_edited_text:
            export_data['texts'] = export_data.get('texts', {})
            export_data['texts']['edited'] = doc_pair.edited_text

        calibration = options.get('calibration')
        if calibration:
            export_data['calibration'] = calibration
        
        return json.dumps(export_data, indent=2)


class PDFExporter:
    """Export to PDF format (using reportlab)."""
    
    @staticmethod
    def export(
        analysis_result,
        doc_pair,
        original_metadata: Dict,
        edited_metadata: Dict,
        **options
    ) -> bytes:
        """
        Generate PDF report using reportlab.
        
        Returns:
            BytesIO object with PDF data
        """
        pdf_content = io.BytesIO()
        doc = SimpleDocTemplate(pdf_content, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=30,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        )
        story.append(Paragraph("VoiceTracer Analysis Report", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Metadata
        story.append(Paragraph(f"<b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Text Statistics
        story.append(Paragraph("Text Statistics", styles['Heading2']))
        stats_data = [
            ['Metric', 'Original', 'Edited', 'Change'],
            ['Word Count', str(original_metadata.get('word_count', 0)), str(edited_metadata.get('word_count', 0)), 
             str(edited_metadata.get('word_count', 0) - original_metadata.get('word_count', 0))],
            ['Character Count', str(original_metadata.get('character_count', 0)), str(edited_metadata.get('character_count', 0)),
             str(edited_metadata.get('character_count', 0) - original_metadata.get('character_count', 0))],
            ['Sentence Count', str(original_metadata.get('sentence_count', 0)), str(edited_metadata.get('sentence_count', 0)),
             str(edited_metadata.get('sentence_count', 0) - original_metadata.get('sentence_count', 0))],
        ]
        stats_table = Table(stats_data, colWidths=[1.5*inch, 1.5*inch, 1.5*inch, 1.5*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Metrics Comparison
        story.append(Paragraph("Metric Comparison", styles['Heading2']))
        metrics_data = [
            ['Metric', 'Original', 'Edited', 'Absolute Shift (Δ)'],
            ['Burstiness', f"{analysis_result.original_metrics.burstiness:.3f}", 
             f"{analysis_result.edited_metrics.burstiness:.3f}",
             f"{analysis_result.metric_deltas.burstiness_delta:+.3f}"],
            ['Lexical Diversity', f"{analysis_result.original_metrics.lexical_diversity:.3f}",
             f"{analysis_result.edited_metrics.lexical_diversity:.3f}",
             f"{analysis_result.metric_deltas.lexical_diversity_delta:+.3f}"],
            ['Syntactic Complexity', f"{analysis_result.original_metrics.syntactic_complexity:.3f}",
             f"{analysis_result.edited_metrics.syntactic_complexity:.3f}",
             f"{analysis_result.metric_deltas.syntactic_complexity_delta:+.3f}"],
            ['AI-ism Likelihood', f"{analysis_result.original_metrics.ai_ism_likelihood:.1f}",
             f"{analysis_result.edited_metrics.ai_ism_likelihood:.1f}",
             f"{analysis_result.metric_deltas.ai_ism_delta:+.1f}"],
              ['Function Word Ratio', f"{analysis_result.original_metrics.function_word_ratio:.3f}",
               f"{analysis_result.edited_metrics.function_word_ratio:.3f}",
               f"{analysis_result.metric_deltas.function_word_ratio_delta:+.3f}"],
              ['Discourse Marker Density', f"{analysis_result.original_metrics.discourse_marker_density:.2f}",
               f"{analysis_result.edited_metrics.discourse_marker_density:.2f}",
               f"{analysis_result.metric_deltas.discourse_marker_density_delta:+.2f}"],
              ['Information Density', f"{analysis_result.original_metrics.information_density:.3f}",
               f"{analysis_result.edited_metrics.information_density:.3f}",
               f"{analysis_result.metric_deltas.information_density_delta:+.3f}"],
              ['Epistemic Hedging', f"{analysis_result.original_metrics.epistemic_hedging:.3f}",
               f"{analysis_result.edited_metrics.epistemic_hedging:.3f}",
               f"{analysis_result.metric_deltas.epistemic_hedging_delta:+.3f}"],
        ]
        metrics_table = Table(metrics_data, colWidths=[1.5*inch, 1.5*inch, 1.5*inch, 1.5*inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(metrics_table)
        story.append(Spacer(1, 0.3*inch))

        verdicts = _build_metric_verdicts(analysis_result, options.get("calibration"))
        verdict_rows = [
            ['Metric', 'Final Verdict'],
            ['Burstiness', verdicts.get('burstiness', 'Authorial voice retained')],
            ['Lexical Diversity', verdicts.get('lexical_diversity', 'Authorial voice retained')],
            ['Syntactic Complexity', verdicts.get('syntactic_complexity', 'Authorial voice retained')],
            ['AI-ism Likelihood', verdicts.get('ai_ism_likelihood', 'Authorial voice retained')],
            ['Function Word Ratio', verdicts.get('function_word_ratio', 'Authorial voice retained')],
            ['Discourse Marker Density', verdicts.get('discourse_marker_density', 'Authorial voice retained')],
            ['Information Density', verdicts.get('information_density', 'Authorial voice retained')],
            ['Epistemic Hedging', verdicts.get('epistemic_hedging', 'Authorial voice retained')],
        ]
        verdict_table = Table(verdict_rows, colWidths=[2.2*inch, 3.8*inch])
        verdict_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#444444')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.whitesmoke),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(Paragraph("Final Verdicts", styles['Heading2']))
        story.append(verdict_table)
        story.append(Spacer(1, 0.3*inch))

        calibration = options.get('calibration')
        if calibration:
            story.append(Paragraph("Calibration Standards", styles['Heading2']))
            calibration_rows = [['Mode', 'Metric', 'Human Standard', 'AI Standard']]
            for mode in ['default', 'adjusted']:
                for metric, human_value in calibration.get(mode, {}).get('human', {}).items():
                    ai_value = calibration.get(mode, {}).get('ai', {}).get(metric, '')
                    calibration_rows.append([mode, str(metric), str(human_value), str(ai_value)])

            calibration_table = Table(calibration_rows, colWidths=[1.2*inch, 1.6*inch, 1.6*inch, 1.6*inch])
            calibration_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#444444')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.whitesmoke),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(calibration_table)
            story.append(Spacer(1, 0.2*inch))

            impact_rows = [['Metric', 'Original Δ', 'Edited Δ']]
            impact = calibration.get('impact', {})
            for metric, orig_delta in impact.get('original', {}).items():
                edit_delta = impact.get('edited', {}).get(metric, '')
                impact_rows.append([str(metric), f"{orig_delta:+.3f}", f"{edit_delta:+.3f}"])

            impact_table = Table(impact_rows, colWidths=[2.0*inch, 2.0*inch, 2.0*inch])
            impact_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(Paragraph("Calibration Impact (Score Δ)", styles['Heading3']))
            story.append(impact_table)
            story.append(Spacer(1, 0.2*inch))

            notes = calibration.get('notes', {})
            if notes:
                story.append(Paragraph("Calibration Notes", styles['Heading3']))
                for text in notes.values():
                    story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 0.2*inch))

        include_charts = options.get("include_charts", False)
        if include_charts:
            story.append(PageBreak())
            story.append(Paragraph("Visualizations", styles['Heading2']))

            chart_errors: List[str] = []

            def add_chart(fig, title, width=6.5*inch, height=4.0*inch):
                story.append(Paragraph(title, styles['Heading3']))
                try:
                    img_bytes = _plotly_to_image(fig)
                    img = Image(io.BytesIO(img_bytes))
                    img.drawWidth = width
                    img.drawHeight = height
                    story.append(img)
                except Exception as exc:
                    error_text = f"{type(exc).__name__}: {str(exc)}"
                    chart_errors.append(f"{title} -> {error_text}")
                    story.append(Paragraph(f"{title} (chart unavailable)", styles['Normal']))
                story.append(Spacer(1, 0.2*inch))

            radar_orig = _normalized_metric_dict(analysis_result.original_metrics)
            radar_edit = _normalized_metric_dict(analysis_result.edited_metrics)
            bar_orig = _raw_metric_dict(analysis_result.original_metrics)
            bar_edit = _raw_metric_dict(analysis_result.edited_metrics)
            deltas_dict = {
                'burstiness_pct_change': analysis_result.metric_deltas.burstiness_pct_change,
                'lexical_diversity_pct_change': analysis_result.metric_deltas.lexical_diversity_pct_change,
                'syntactic_complexity_pct_change': analysis_result.metric_deltas.syntactic_complexity_pct_change,
                'ai_ism_likelihood_pct_change': analysis_result.metric_deltas.ai_ism_pct_change,
                'function_word_ratio_pct_change': analysis_result.metric_deltas.function_word_ratio_pct_change,
                'discourse_marker_density_pct_change': analysis_result.metric_deltas.discourse_marker_density_pct_change,
                'information_density_pct_change': analysis_result.metric_deltas.information_density_pct_change,
                'epistemic_hedging_pct_change': analysis_result.metric_deltas.epistemic_hedging_pct_change,
            }

            add_chart(
                RadarChartGenerator.create_metric_radar(radar_orig, radar_edit),
                "8-Axis Radar Chart",
            )
            add_chart(
                BarChartGenerator.create_metric_comparison(bar_orig, bar_edit),
                "Metric Comparison (Bar Chart)",
            )
            add_chart(
                DeltaVisualization.create_delta_chart(deltas_dict),
                "Metric Shifts (Delta Chart)",
            )

            burst_bar_fig, _ = BurstinessVisualization.create_sentence_length_bars(
                doc_pair.original_text,
                doc_pair.edited_text,
            )
            add_chart(burst_bar_fig, "Burstiness: Word Count per Sentence", height=3.6*inch)
            add_chart(
                BurstinessVisualization.create_fluctuation_curve(
                    doc_pair.original_text,
                    doc_pair.edited_text,
                ),
                "Burstiness: Fluctuation Pattern",
                height=3.6*inch,
            )

            story.append(PageBreak())
            story.append(Paragraph("Individual Metric Charts", styles['Heading2']))
            panels = IndividualMetricCharts.create_metric_panels(bar_orig, bar_edit)
            for name, fig in panels:
                add_chart(fig, name, height=3.0*inch)

            if chart_errors:
                story.append(PageBreak())
                story.append(Paragraph("Visualization Export Errors", styles['Heading2']))
                for text in chart_errors:
                    story.append(Paragraph(text, styles['Normal']))
        
        # AI-isms detected
        if analysis_result.ai_isms:
            story.append(PageBreak())
            story.append(Paragraph("AI-isms Detected", styles['Heading2']))
            for ai_ism in analysis_result.ai_isms[:5]:
                story.append(Paragraph(f"<b>{ai_ism.get('phrase', 'N/A')}</b> ({ai_ism.get('category', 'N/A')})", styles['Normal']))
                story.append(Paragraph(f"<i>{ai_ism.get('context', 'N/A')[:100]}</i>", styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        pdf_content.seek(0)
        return pdf_content.getvalue()


class ExcelExporter:
    """Export to Excel (XLSX) format."""
    
    @staticmethod
    def export(
        analysis_result,
        doc_pair,
        original_metadata: Dict,
        edited_metadata: Dict,
        **options
    ) -> BinaryIO:
        """
        Generate Excel workbook with multiple sheets.
        
        Sheets:
        1. Summary - Key metrics and stats
        2. Detailed Metrics - Full metric values
        3. Text Statistics - Word count, etc.
        4. AI-isms - Detected phrases and context
        5. Recommendations - Actionable suggestions
        
        Returns:
            BytesIO object with XLSX data
        """
        output = io.BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Sheet 1: Summary
            summary_data = {
                'Metric': [
                    'Burstiness',
                    'Lexical Diversity',
                    'Syntactic Complexity',
                    'AI-ism Likelihood',
                    'Function Word Ratio',
                    'Discourse Marker Density',
                    'Information Density',
                    'Epistemic Hedging',
                ],
                'Original': [
                    analysis_result.original_metrics.burstiness,
                    analysis_result.original_metrics.lexical_diversity,
                    analysis_result.original_metrics.syntactic_complexity,
                    analysis_result.original_metrics.ai_ism_likelihood,
                    analysis_result.original_metrics.function_word_ratio,
                    analysis_result.original_metrics.discourse_marker_density,
                    analysis_result.original_metrics.information_density,
                    analysis_result.original_metrics.epistemic_hedging,
                ],
                'Edited': [
                    analysis_result.edited_metrics.burstiness,
                    analysis_result.edited_metrics.lexical_diversity,
                    analysis_result.edited_metrics.syntactic_complexity,
                    analysis_result.edited_metrics.ai_ism_likelihood,
                    analysis_result.edited_metrics.function_word_ratio,
                    analysis_result.edited_metrics.discourse_marker_density,
                    analysis_result.edited_metrics.information_density,
                    analysis_result.edited_metrics.epistemic_hedging,
                ],
                'Absolute Shift (Δ)': [
                    analysis_result.metric_deltas.burstiness_delta,
                    analysis_result.metric_deltas.lexical_diversity_delta,
                    analysis_result.metric_deltas.syntactic_complexity_delta,
                    analysis_result.metric_deltas.ai_ism_delta,
                    analysis_result.metric_deltas.function_word_ratio_delta,
                    analysis_result.metric_deltas.discourse_marker_density_delta,
                    analysis_result.metric_deltas.information_density_delta,
                    analysis_result.metric_deltas.epistemic_hedging_delta,
                ],
            }
            df_summary = pd.DataFrame(summary_data)
            df_summary.to_excel(writer, sheet_name='Summary', index=False)
            
            # Sheet 2: Text Statistics
            stats_data = {
                'Statistic': ['Word Count', 'Character Count', 'Sentence Count', 'Avg Sentence Length'],
                'Original': [
                    original_metadata.get('word_count', 0),
                    original_metadata.get('char_count', 0),
                    original_metadata.get('sentence_count', 0),
                    original_metadata.get('avg_sentence_length', 0),
                ],
                'Edited': [
                    edited_metadata.get('word_count', 0),
                    edited_metadata.get('char_count', 0),
                    edited_metadata.get('sentence_count', 0),
                    edited_metadata.get('avg_sentence_length', 0),
                ],
            }
            df_stats = pd.DataFrame(stats_data)
            df_stats.to_excel(writer, sheet_name='Statistics', index=False)

            calibration = options.get('calibration')
            if calibration:
                standards_rows = []
                for mode in ['default', 'adjusted']:
                    for metric, human_value in calibration.get(mode, {}).get('human', {}).items():
                        ai_value = calibration.get(mode, {}).get('ai', {}).get(metric, '')
                        standards_rows.append({
                            'Mode': mode,
                            'Metric': metric,
                            'Human Standard': human_value,
                            'AI Standard': ai_value,
                        })

                impact_rows = []
                impact = calibration.get('impact', {})
                for metric, orig_delta in impact.get('original', {}).items():
                    impact_rows.append({
                        'Metric': metric,
                        'Original Δ': orig_delta,
                        'Edited Δ': impact.get('edited', {}).get(metric, ''),
                    })

                notes_rows = [
                    {'Note': text}
                    for text in calibration.get('notes', {}).values()
                ]

                df_standards = pd.DataFrame(standards_rows)
                df_impact = pd.DataFrame(impact_rows)
                df_notes = pd.DataFrame(notes_rows)

                df_standards.to_excel(writer, sheet_name='Calibration Standards', index=False)
                df_impact.to_excel(writer, sheet_name='Calibration Impact', index=False)
                if not df_notes.empty:
                    df_notes.to_excel(writer, sheet_name='Calibration Notes', index=False)
        
        output.seek(0)
        return output


class DocxExporter:
    """Export to Word (DOCX) format."""
    
    @staticmethod
    def export(
        analysis_result,
        doc_pair,
        original_metadata: Dict,
        edited_metadata: Dict,
        **options
    ) -> bytes:
        """
        Generate Word document with analysis using python-docx.
        
        Returns:
            BytesIO object with DOCX data
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('VoiceTracer Analysis Report', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        # Text Statistics
        doc.add_heading('Text Statistics', 1)
        stats_table = doc.add_table(rows=4, cols=4)
        stats_table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = stats_table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Original'
        header_cells[2].text = 'Edited'
        header_cells[3].text = 'Change'
        
        # Data rows
        rows_data = [
            ['Word Count', str(original_metadata.get('word_count', 0)), str(edited_metadata.get('word_count', 0)),
             str(edited_metadata.get('word_count', 0) - original_metadata.get('word_count', 0))],
            ['Character Count', str(original_metadata.get('character_count', 0)), str(edited_metadata.get('character_count', 0)),
             str(edited_metadata.get('character_count', 0) - original_metadata.get('character_count', 0))],
            ['Sentence Count', str(original_metadata.get('sentence_count', 0)), str(edited_metadata.get('sentence_count', 0)),
             str(edited_metadata.get('sentence_count', 0) - original_metadata.get('sentence_count', 0))],
        ]
        
        for i, row_data in enumerate(rows_data, start=1):
            cells = stats_table.rows[i].cells
            for j, text in enumerate(row_data):
                cells[j].text = text
        
        doc.add_paragraph()
        
        # Metrics Comparison
        doc.add_heading('Metric Comparison', 1)
        metrics_table = doc.add_table(rows=9, cols=4)
        metrics_table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = metrics_table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Original'
        header_cells[2].text = 'Edited'
        header_cells[3].text = 'Absolute Shift (Δ)'
        
        # Metrics data
        metrics_rows = [
            ['Burstiness', f"{analysis_result.original_metrics.burstiness:.3f}", 
             f"{analysis_result.edited_metrics.burstiness:.3f}",
             f"{analysis_result.metric_deltas.burstiness_delta:+.3f}"],
            ['Lexical Diversity', f"{analysis_result.original_metrics.lexical_diversity:.3f}",
             f"{analysis_result.edited_metrics.lexical_diversity:.3f}",
             f"{analysis_result.metric_deltas.lexical_diversity_delta:+.3f}"],
            ['Syntactic Complexity', f"{analysis_result.original_metrics.syntactic_complexity:.3f}",
             f"{analysis_result.edited_metrics.syntactic_complexity:.3f}",
             f"{analysis_result.metric_deltas.syntactic_complexity_delta:+.3f}"],
            ['AI-ism Likelihood', f"{analysis_result.original_metrics.ai_ism_likelihood:.1f}",
             f"{analysis_result.edited_metrics.ai_ism_likelihood:.1f}",
             f"{analysis_result.metric_deltas.ai_ism_delta:+.1f}"],
              ['Function Word Ratio', f"{analysis_result.original_metrics.function_word_ratio:.3f}",
               f"{analysis_result.edited_metrics.function_word_ratio:.3f}",
               f"{analysis_result.metric_deltas.function_word_ratio_delta:+.3f}"],
              ['Discourse Marker Density', f"{analysis_result.original_metrics.discourse_marker_density:.2f}",
               f"{analysis_result.edited_metrics.discourse_marker_density:.2f}",
               f"{analysis_result.metric_deltas.discourse_marker_density_delta:+.2f}"],
              ['Information Density', f"{analysis_result.original_metrics.information_density:.3f}",
               f"{analysis_result.edited_metrics.information_density:.3f}",
               f"{analysis_result.metric_deltas.information_density_delta:+.3f}"],
              ['Epistemic Hedging', f"{analysis_result.original_metrics.epistemic_hedging:.3f}",
               f"{analysis_result.edited_metrics.epistemic_hedging:.3f}",
               f"{analysis_result.metric_deltas.epistemic_hedging_delta:+.3f}"],
        ]
        
        for i, row_data in enumerate(metrics_rows, start=1):
            cells = metrics_table.rows[i].cells
            for j, text in enumerate(row_data):
                cells[j].text = text
        
        doc.add_paragraph()

        calibration = options.get('calibration')
        if calibration:
            doc.add_heading('Calibration Standards', 1)
            standards_table = doc.add_table(rows=1, cols=4)
            standards_table.style = 'Light Grid Accent 1'
            header_cells = standards_table.rows[0].cells
            header_cells[0].text = 'Mode'
            header_cells[1].text = 'Metric'
            header_cells[2].text = 'Human Standard'
            header_cells[3].text = 'AI Standard'

            for mode in ['default', 'adjusted']:
                for metric, human_value in calibration.get(mode, {}).get('human', {}).items():
                    ai_value = calibration.get(mode, {}).get('ai', {}).get(metric, '')
                    row_cells = standards_table.add_row().cells
                    row_cells[0].text = mode
                    row_cells[1].text = str(metric)
                    row_cells[2].text = str(human_value)
                    row_cells[3].text = str(ai_value)

            doc.add_paragraph()
            doc.add_heading('Calibration Impact (Score Δ)', 2)
            impact_table = doc.add_table(rows=1, cols=3)
            impact_table.style = 'Light Grid Accent 1'
            header_cells = impact_table.rows[0].cells
            header_cells[0].text = 'Metric'
            header_cells[1].text = 'Original Δ'
            header_cells[2].text = 'Edited Δ'

            impact = calibration.get('impact', {})
            for metric, orig_delta in impact.get('original', {}).items():
                row_cells = impact_table.add_row().cells
                row_cells[0].text = str(metric)
                row_cells[1].text = f"{orig_delta:+.3f}"
                row_cells[2].text = f"{impact.get('edited', {}).get(metric, 0.0):+.3f}"

            notes = calibration.get('notes', {})
            if notes:
                doc.add_paragraph()
                doc.add_heading('Calibration Notes', 2)
                for text in notes.values():
                    doc.add_paragraph(text)

            doc.add_paragraph()
        
        # AI-isms Detected
        if analysis_result.ai_isms:
            doc.add_heading('AI-isms Detected', 1)
            for ai_ism in analysis_result.ai_isms[:5]:
                phrase = ai_ism.get('phrase', 'N/A')
                category = ai_ism.get('category', 'N/A')
                context = ai_ism.get('context', 'N/A')[:150]
                
                p = doc.add_paragraph()
                p.add_run(f"{phrase} ").bold = True
                p.add_run(f"({category})")
                
                doc.add_paragraph(context, style='List Bullet')
            
            doc.add_paragraph()
        
        # Original Text (if included)
        if options.get('include_original_text', False):
            doc.add_page_break()
            doc.add_heading('Original Text', 1)
            doc.add_paragraph(doc_pair.original_text)
        
        # Edited Text (if included)
        if options.get('include_edited_text', False):
            doc.add_page_break()
            doc.add_heading('Edited Text', 1)
            doc.add_paragraph(doc_pair.edited_text)
        
        # Save to BytesIO
        docx_content = io.BytesIO()
        doc.save(docx_content)
        docx_content.seek(0)
        return docx_content.getvalue()


class PowerPointExporter:
    """Export to PowerPoint (PPTX) format."""
    
    @staticmethod
    def export(
        analysis_result,
        doc_pair,
        original_metadata: Dict,
        edited_metadata: Dict,
        **options
    ) -> BinaryIO:
        """
        Generate PowerPoint presentation with slides.
        
        Note: Uses python-pptx to create presentations with:
        - 9 slides covering analysis
        - Embedded charts
        - Speaker notes
        - Editable by instructors
        
        Returns:
            BytesIO object with PPTX data
        """
        # Placeholder - full implementation would use python-pptx
        pptx_content = io.BytesIO()
        pptx_content.write(b'PPTX Export - Not yet implemented')
        pptx_content.seek(0)
        return pptx_content


class ExportFactory:
    """Factory for creating exporters based on format."""
    
    EXPORTERS = {
        'csv': CSVExporter,
        'json': JSONExporter,
        'xlsx': ExcelExporter,
        'docx': DocxExporter,
        'pptx': PowerPointExporter,
        'pdf': PDFExporter,
    }
    
    @staticmethod
    def export(
        format_type: str,
        analysis_result,
        doc_pair,
        original_metadata: Dict,
        edited_metadata: Dict,
        **options
    ) -> Any:
        """
        Generate export in specified format.
        
        Args:
            format_type: 'csv', 'json', 'xlsx', 'docx', 'pptx', 'pdf'
            analysis_result: AnalysisResult object
            doc_pair: DocumentPair object
            original_metadata: Original text metadata dict
            edited_metadata: Edited text metadata dict
            **options: Additional format-specific options
        
        Returns:
            Export data (string for CSV/JSON, BytesIO for binary formats)
        """
        if format_type not in ExportFactory.EXPORTERS:
            raise ValueError(f"Unsupported format: {format_type}")
        
        exporter_class = ExportFactory.EXPORTERS[format_type]
        return exporter_class.export(
            analysis_result,
            doc_pair,
            original_metadata,
            edited_metadata,
            **options
        )
